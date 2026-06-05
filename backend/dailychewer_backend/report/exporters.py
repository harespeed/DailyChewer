"""Export weekly reports to markdown, docx, xlsx, and csv."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from dailychewer_backend.config import SUPPORTED_EXPORT_FORMATS, WEEKDAY_ORDER
from dailychewer_backend.models import DailyReport, MonthlyReport, WeeklyReport


def export_weekly_report(report: WeeklyReport, output_path: Path, fmt: str) -> Path:
    """Export a weekly report to the requested file format."""

    normalized_fmt = fmt.lower()
    if normalized_fmt not in SUPPORTED_EXPORT_FORMATS:
        raise ValueError(f"不支持的导出格式：{fmt}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    if normalized_fmt == "markdown":
        output_path.write_text(render_weekly_markdown(report), encoding="utf-8")
        return output_path
    if normalized_fmt == "docx":
        export_weekly_docx(report, output_path)
        return output_path
    if normalized_fmt == "xlsx":
        export_weekly_xlsx(report, output_path)
        return output_path
    if normalized_fmt == "csv":
        export_weekly_csv(report, output_path)
        return output_path
    raise ValueError(f"不支持的导出格式：{fmt}")


def export_monthly_report(report: MonthlyReport, output_path: Path, fmt: str) -> Path:
    """Export a monthly report to the requested file format."""

    normalized_fmt = fmt.lower()
    if normalized_fmt not in SUPPORTED_EXPORT_FORMATS:
        raise ValueError(f"不支持的导出格式：{fmt}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    if normalized_fmt == "markdown":
        output_path.write_text(render_monthly_markdown(report), encoding="utf-8")
        return output_path
    if normalized_fmt == "docx":
        export_monthly_docx(report, output_path)
        return output_path
    if normalized_fmt == "xlsx":
        export_monthly_xlsx(report, output_path)
        return output_path
    if normalized_fmt == "csv":
        export_monthly_csv(report, output_path)
        return output_path
    raise ValueError(f"不支持的导出格式：{fmt}")


def render_weekly_markdown(report: WeeklyReport) -> str:
    """Render a weekly report to markdown text."""

    parts = [
        "# DailyChewer 周报",
        "",
        "## 周报周期",
        "",
        f"{report.start_date} 至 {report.end_date}",
        "",
        "---",
    ]

    for daily_report in _iter_report_days(report):
        parts.extend(
            [
                "",
                f"## {daily_report.weekday} {daily_report.date}",
                "",
                "### 上午",
                "",
                *_render_markdown_section(daily_report.morning),
                "",
                "### 下午",
                "",
                *_render_markdown_section(daily_report.afternoon),
                "",
                "---",
            ]
        )

    parts.extend(
        [
            "",
            "## 本周收获",
            "",
            *_render_markdown_bullets(report.weekly_gains or ["暂无可总结收获"]),
            "",
        ]
    )
    return "\n".join(parts).strip() + "\n"


def render_monthly_markdown(report: MonthlyReport) -> str:
    """Render a monthly report to markdown text."""

    sections = [
        ("本月主要工作", report.main_work),
        ("关键进展", report.key_progress),
        ("问题与解决方案", report.problems_and_solutions),
        ("个人成长", report.personal_growth),
        ("本月收获", report.monthly_gains),
        ("下月改进方向", report.next_improvements),
    ]
    parts = [
        "# DailyChewer 月报",
        "",
        "## 月报周期",
        "",
        report.month,
        "",
    ]
    for title, items in sections:
        parts.extend(
            [
                f"## {title}",
                "",
                *_render_markdown_bullets(items or ["暂无可总结内容"]),
                "",
            ]
        )
    return "\n".join(parts).strip() + "\n"


def export_weekly_docx(report: WeeklyReport, output_path: Path) -> None:
    """Export a weekly report into a DOCX document."""

    from docx import Document

    document = Document()
    document.add_heading("DailyChewer 周报", level=1)
    document.add_heading("周报周期", level=2)
    document.add_paragraph(f"{report.start_date} 至 {report.end_date}")

    for daily_report in _iter_report_days(report):
        document.add_heading(f"{daily_report.weekday} {daily_report.date}", level=2)
        document.add_heading("上午", level=3)
        _append_docx_section(document, daily_report.morning)
        document.add_heading("下午", level=3)
        _append_docx_section(document, daily_report.afternoon)

    document.add_heading("本周收获", level=2)
    for item in report.weekly_gains or ["暂无可总结收获"]:
        document.add_paragraph(item, style="List Bullet")

    document.save(output_path)


def export_monthly_docx(report: MonthlyReport, output_path: Path) -> None:
    """Export a monthly report into a DOCX document."""

    from docx import Document

    document = Document()
    document.add_heading("DailyChewer 月报", level=1)
    document.add_heading("月报周期", level=2)
    document.add_paragraph(report.month)
    for title, items in [
        ("本月主要工作", report.main_work),
        ("关键进展", report.key_progress),
        ("问题与解决方案", report.problems_and_solutions),
        ("个人成长", report.personal_growth),
        ("本月收获", report.monthly_gains),
        ("下月改进方向", report.next_improvements),
    ]:
        document.add_heading(title, level=2)
        for item in items or ["暂无可总结内容"]:
            document.add_paragraph(item, style="List Bullet")
    document.save(output_path)


def export_weekly_xlsx(report: WeeklyReport, output_path: Path) -> None:
    """Export a weekly report into a workbook with detail and gains sheets."""

    from openpyxl import Workbook

    workbook = Workbook()
    detail_sheet = workbook.active
    detail_sheet.title = "Weekly Report"
    headers = [
        "Week",
        "Date",
        "Weekday",
        "Period",
        "Work Content",
        "Personal Growth",
        "Problems",
        "Solutions",
    ]
    detail_sheet.append(headers)
    for row in _iter_tabular_rows(report):
        detail_sheet.append(row)

    gains_sheet = workbook.create_sheet("Weekly Gains")
    gains_sheet.append(["Week", "Gain"])
    for item in report.weekly_gains or ["暂无可总结收获"]:
        gains_sheet.append([report.week, item])

    workbook.save(output_path)


def export_monthly_xlsx(report: MonthlyReport, output_path: Path) -> None:
    """Export a monthly report into a workbook."""

    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Monthly Report"
    sheet.append(["Month", "Section", "Content"])
    for title, items in [
        ("Main Work", report.main_work),
        ("Key Progress", report.key_progress),
        ("Problems and Solutions", report.problems_and_solutions),
        ("Personal Growth", report.personal_growth),
        ("Monthly Gains", report.monthly_gains),
        ("Next Improvements", report.next_improvements),
    ]:
        for item in items or ["暂无可总结内容"]:
            sheet.append([report.month, title, item])
    workbook.save(output_path)


def export_weekly_csv(report: WeeklyReport, output_path: Path) -> None:
    """Export a weekly report into a flat CSV file."""

    import pandas as pd

    rows = [
        {
            "Type": "daily",
            "Week": row[0],
            "Date": row[1],
            "Weekday": row[2],
            "Period": row[3],
            "Work Content": row[4],
            "Personal Growth": row[5],
            "Problems": row[6],
            "Solutions": row[7],
            "Content": "",
        }
        for row in _iter_tabular_rows(report)
    ]
    for item in report.weekly_gains or ["暂无可总结收获"]:
        rows.append(
            {
                "Type": "weekly_gain",
                "Week": report.week,
                "Date": "",
                "Weekday": "",
                "Period": "",
                "Work Content": "",
                "Personal Growth": "",
                "Problems": "",
                "Solutions": "",
                "Content": item,
            }
        )
    dataframe = pd.DataFrame(rows)
    dataframe.to_csv(output_path, index=False, encoding="utf-8-sig")


def export_monthly_csv(report: MonthlyReport, output_path: Path) -> None:
    """Export a monthly report into CSV rows."""

    import pandas as pd

    rows = []
    for title, items in [
        ("Main Work", report.main_work),
        ("Key Progress", report.key_progress),
        ("Problems and Solutions", report.problems_and_solutions),
        ("Personal Growth", report.personal_growth),
        ("Monthly Gains", report.monthly_gains),
        ("Next Improvements", report.next_improvements),
    ]:
        for item in items or ["暂无可总结内容"]:
            rows.append({"Month": report.month, "Section": title, "Content": item})
    pd.DataFrame(rows).to_csv(output_path, index=False, encoding="utf-8-sig")


def _render_markdown_section(section: Any) -> list[str]:
    """Render one half-day section in markdown."""

    return [
        "#### 工作内容",
        *_render_markdown_bullets(section.work_content),
        "",
        "#### 个人成长",
        *_render_markdown_bullets(section.personal_growth),
        "",
        "#### 问题总结",
        *_render_markdown_bullets(section.problems),
        "",
        "#### 解决方案",
        *_render_markdown_bullets(section.solutions),
    ]


def _render_markdown_bullets(items: list[str]) -> list[str]:
    """Render a list of strings into markdown bullet lines."""

    return [f"- {item}" for item in (items or ["暂无日报记录"])]


def _append_docx_section(document: Any, section: Any) -> None:
    """Append section headings and bullet items to a DOCX document."""

    mapping = [
        ("工作内容", section.work_content),
        ("个人成长", section.personal_growth),
        ("问题总结", section.problems),
        ("解决方案", section.solutions),
    ]
    for title, items in mapping:
        document.add_paragraph(title)
        for item in items or ["暂无日报记录"]:
            document.add_paragraph(item, style="List Bullet")


def _iter_tabular_rows(report: WeeklyReport) -> list[list[str]]:
    """Flatten the weekly report into rows for CSV/XLSX export."""

    rows: list[list[str]] = []
    for daily_report in _iter_report_days(report):
        rows.extend(
            [
                _section_to_row(report.week, daily_report, "Morning", daily_report.morning),
                _section_to_row(report.week, daily_report, "Afternoon", daily_report.afternoon),
            ]
        )
    return rows


def _iter_report_days(report: WeeklyReport) -> list[DailyReport]:
    """Return weekly report days in display order for standard or custom ranges."""

    if report.date_range is None and all(day in report.days for day in WEEKDAY_ORDER):
        return [report.days[day] for day in WEEKDAY_ORDER]
    return sorted(report.days.values(), key=lambda item: item.date)


def _section_to_row(week: str, daily_report: DailyReport, period: str, section: Any) -> list[str]:
    """Convert one half-day section into a single tabular row."""

    return [
        week,
        daily_report.date,
        daily_report.weekday,
        period,
        "\n".join(section.work_content or ["暂无日报记录"]),
        "\n".join(section.personal_growth or ["暂无日报记录"]),
        "\n".join(section.problems or ["暂无日报记录"]),
        "\n".join(section.solutions or ["暂无日报记录"]),
    ]
