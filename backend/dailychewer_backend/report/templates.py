"""Daily report template generation helpers."""

from __future__ import annotations

from pathlib import Path


def generate_daily_template(output_path: Path, date_str: str, fmt: str) -> Path:
    """Generate a daily report template in markdown, csv, xlsx, or docx format."""

    normalized_fmt = fmt.lower()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if normalized_fmt == "markdown":
        output_path.write_text(_markdown_template(date_str), encoding="utf-8")
        return output_path
    if normalized_fmt == "csv":
        _write_csv_template(output_path, date_str)
        return output_path
    if normalized_fmt == "xlsx":
        _write_xlsx_template(output_path, date_str)
        return output_path
    if normalized_fmt == "docx":
        _write_docx_template(output_path, date_str)
        return output_path
    raise ValueError(f"Unsupported template format: {fmt}")


def _markdown_template(date_str: str) -> str:
    return f"""# {date_str} 日报

## 上午

### 工作内容
- 

### 个人成长
- 

### 问题总结
- 

### 解决方案
- 

## 下午

### 工作内容
- 

### 个人成长
- 

### 问题总结
- 

### 解决方案
- 

## 备注
- 
"""


def _write_csv_template(output_path: Path, date_str: str) -> None:
    import csv

    rows = [
        ["Date", "Period", "Category", "Content"],
        [date_str, "Morning", "Work Content", ""],
        [date_str, "Morning", "Personal Growth", ""],
        [date_str, "Morning", "Problems", ""],
        [date_str, "Morning", "Solutions", ""],
        [date_str, "Afternoon", "Work Content", ""],
        [date_str, "Afternoon", "Personal Growth", ""],
        [date_str, "Afternoon", "Problems", ""],
        [date_str, "Afternoon", "Solutions", ""],
    ]
    with output_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerows(rows)


def _write_xlsx_template(output_path: Path, date_str: str) -> None:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Daily Template"
    for row in [
        ["Date", "Period", "Category", "Content"],
        [date_str, "Morning", "Work Content", ""],
        [date_str, "Morning", "Personal Growth", ""],
        [date_str, "Morning", "Problems", ""],
        [date_str, "Morning", "Solutions", ""],
        [date_str, "Afternoon", "Work Content", ""],
        [date_str, "Afternoon", "Personal Growth", ""],
        [date_str, "Afternoon", "Problems", ""],
        [date_str, "Afternoon", "Solutions", ""],
    ]:
        sheet.append(row)
    workbook.save(output_path)


def _write_docx_template(output_path: Path, date_str: str) -> None:
    from docx import Document

    document = Document()
    document.add_heading(f"{date_str} 日报", level=1)
    for period in ["上午", "下午"]:
        document.add_heading(period, level=2)
        for title in ["工作内容", "个人成长", "问题总结", "解决方案"]:
            document.add_heading(title, level=3)
            document.add_paragraph("-", style="List Bullet")
    document.add_heading("备注", level=2)
    document.add_paragraph("-", style="List Bullet")
    document.save(output_path)
