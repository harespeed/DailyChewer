"""DOCX report parser."""

from __future__ import annotations

from pathlib import Path

from dailychewer.parser.base import BaseParser
from dailychewer.utils.text_utils import ensure_non_empty_text, rows_to_markdown_table


class DocxParser(BaseParser):
    """Read paragraphs and tables from a DOCX report."""

    def parse(self, file_path: Path) -> str:
        """Extract paragraphs and render tables as markdown-style text."""

        from docx import Document

        document = Document(file_path)
        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for index, table in enumerate(document.tables, start=1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            headers = rows[0]
            body = rows[1:] if len(rows) > 1 else []
            table_text = rows_to_markdown_table(headers, body)
            parts.append(f"# Table {index}\n\n{table_text}")

        return ensure_non_empty_text("\n\n".join(parts))
