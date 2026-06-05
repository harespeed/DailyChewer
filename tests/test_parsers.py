from pathlib import Path

import pytest

pytest.importorskip("pandas")
pytest.importorskip("openpyxl")
pytest.importorskip("docx")

from docx import Document
from openpyxl import Workbook

from dailychewer.parser.csv_parser import CSVParser
from dailychewer.parser.docx_parser import DocxParser
from dailychewer.parser.markdown_parser import MarkdownParser
from dailychewer.parser.xlsx_parser import XlsxParser


def test_markdown_parser(tmp_path: Path) -> None:
    file_path = tmp_path / "report.md"
    file_path.write_text("# 2026-06-03\n\n- 完成接口梳理", encoding="utf-8")

    content = MarkdownParser().parse(file_path)

    assert "完成接口梳理" in content


def test_csv_parser(tmp_path: Path) -> None:
    file_path = tmp_path / "report.csv"
    file_path.write_text("task,status\n梳理接口,进行中\n", encoding="utf-8")

    content = CSVParser().parse(file_path)

    assert "| task | status |" in content
    assert "梳理接口" in content


def test_xlsx_parser(tmp_path: Path) -> None:
    file_path = tmp_path / "report.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "日报"
    sheet.append(["task", "status"])
    sheet.append(["整理字段", "完成"])
    workbook.save(file_path)

    content = XlsxParser().parse(file_path)

    assert "# Sheet: 日报" in content
    assert "整理字段" in content


def test_docx_parser(tmp_path: Path) -> None:
    file_path = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("上午梳理需求")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "任务"
    table.rows[0].cells[1].text = "状态"
    table.rows[1].cells[0].text = "接口联调"
    table.rows[1].cells[1].text = "处理中"
    document.save(file_path)

    content = DocxParser().parse(file_path)

    assert "上午梳理需求" in content
    assert "# Table 1" in content
    assert "接口联调" in content
